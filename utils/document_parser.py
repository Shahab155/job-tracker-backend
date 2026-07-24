import io
import pypdf
import docx
from fastapi import HTTPException

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF file bytes using pypdf."""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        extracted = "\n\n".join(text_parts).strip()
        if not extracted:
            raise ValueError("No readable text found in PDF file.")
        return extracted
    except Exception as e:
        if isinstance(e, ValueError):
            raise e
        raise ValueError(f"Failed to parse PDF document: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX file bytes using python-docx."""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)
        extracted = "\n".join(full_text).strip()
        if not extracted:
            raise ValueError("No readable text found in Word document.")
        return extracted
    except Exception as e:
        if isinstance(e, ValueError):
            raise e
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")


def parse_resume_file(file_bytes: bytes, filename: str) -> str:
    """
    Parses resume document bytes based on file extension (.pdf, .docx, .txt).
    Returns cleaned extracted string content.
    """
    fname_lower = filename.lower()
    if fname_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif fname_lower.endswith(".docx") or fname_lower.endswith(".doc"):
        if fname_lower.endswith(".doc") and not fname_lower.endswith(".docx"):
            raise HTTPException(
                status_code=400,
                detail="Legacy .doc format is not supported. Please convert your file to .docx or .pdf."
            )
        return extract_text_from_docx(file_bytes)
    elif fname_lower.endswith(".txt") or fname_lower.endswith(".md"):
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1")
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{filename}'. Only PDF (.pdf), Word (.docx), and Text (.txt) files are supported."
        )
